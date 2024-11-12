import modal

# Define your Modal app

# Model configuration
MODEL_NAME = "Meta-Llama-3.1-8B-Instruct"
MODEL_FILE = "Meta-Llama-3.1-8B-Instruct-Q5_K_M.gguf"
REVISION = "9a8dec50f04fa8fad1dc1e7bc20a84a512e2bb01"
LLAMA_CPP_RELEASE = "b3472"

# Define the image with all dependencies
image = (
    modal.Image.debian_slim(python_version="3.11")
    .pip_install(
        [
            "flask",
            "flask-limiter",
            "opencv-python-headless",
            "pandas",
            "pdf2image",
            "Pillow",
            "python-docx",
            "pytesseract",
            "werkzeug",
            "PyPDF2",
        ]
    )
    .apt_install(["curl",
        "unzip",
        "libgl1",  # Install libGL.so.1
        "poppler-utils",  # Required by pdf2image
        "tesseract-ocr",  # Required by pytesseract
        "libglib2.0-0",   # Additional dependencies
        "libsm6",
        "libxext6",
        "libxrender-dev",])
    .run_commands(
        [
            # Download and unzip llama.cpp binaries
            f"curl -L -O https://github.com/ggerganov/llama.cpp/releases/download/{LLAMA_CPP_RELEASE}/llama-{LLAMA_CPP_RELEASE}-bin-ubuntu-x64.zip",
            f"unzip llama-{LLAMA_CPP_RELEASE}-bin-ubuntu-x64.zip",
            f"apt-get update && apt-get install -y --no-install-recommends \
            curl \
            unzip \
            libgl1 \
            poppler-utils \
            tesseract-ocr \
            libglib2.0-0 \
            libsm6 \
            libxext6 \
            libxrender-dev && \
            apt-get clean && \
            rm -rf /var/lib/apt/lists/*",
        ]
    )
)

image = image.run_commands(
    f"curl -L -O https://huggingface.co/bartowski/{MODEL_NAME}-GGUF/resolve/{REVISION}/{MODEL_FILE}?download=true"
)

app = modal.App("classifier-app")

# Include your source code in the Modal image
mounts = [
    modal.Mount.from_local_dir("src", remote_path="/src"),
    # Add other mounts if necessary
]


# Define the Modal function to run the Llama 3 model
@app.cls(
    allow_concurrent_inputs=15,
    container_idle_timeout=60 * 10,
    timeout=60 * 60,
    image=image,
)
class Model:
    @modal.method()
    def classify_with_llama(self, prompt: str):
        import subprocess
        num_output_tokens = 128
        result = subprocess.run(
            [
                "/build/bin/llama-cli",
                "-m",
                f"/{MODEL_FILE}",
                "-n",
                str(num_output_tokens),
                "-p",
                prompt,
            ],
            capture_output=True,
            text=True,
        )

        print("Llama model output:" + result.stdout)

        # Check for errors
        if result.returncode != 0:
            print("Error running Llama model:")
            print(result.stderr)
            return "Error running Llama model"
        
        response = result.stdout.strip()
        return response

# Define the Flask app within the Modal app
@app.function(
    image=image,
    mounts=mounts,
)
@modal.wsgi_app()
def flask_app():
    from flask import Flask, request, jsonify
    from src.ocr import process_file
    from src.classifier import classify_file_by_name
    from src.ai_assistants import classify_with_llama
    from src.constants import ALLOWED_EXTENSIONS, OUTPUT_DIR, POSSIBLE_FILE_TYPES
    import os
    from PyPDF2 import PdfReader, PdfWriter
    from docx import Document
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address

    app = Flask(__name__)

    limiter = Limiter(
        get_remote_address,
        app=app,
        default_limits=["100 per day", "10 per hour"],
        storage_uri="memory://",
    )

    LARGE_PAGE_COUNT = 5  # Threshold for page count

    def allowed_file(filename):
        return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def trim_pdf_to_first_five_pages(input_path, output_path):
        """Trim the PDF to the first five pages if it has more than five."""
        reader = PdfReader(input_path)
        writer = PdfWriter()

        # Copy only the first five pages if there are more
        for page_num in range(min(LARGE_PAGE_COUNT, len(reader.pages))):
            writer.add_page(reader.pages[page_num])

        with open(output_path, 'wb') as output_file:
            writer.write(output_file)

    def trim_docx_to_first_five_paragraphs(input_path, output_path):
        """Trim the DOCX to the first five paragraphs if it has more than five."""
        document = Document(input_path)
        new_doc = Document()

        # Copy only the first five paragraphs
        for paragraph in document.paragraphs[:LARGE_PAGE_COUNT]:
            new_doc.add_paragraph(paragraph.text)

        new_doc.save(output_path)

    @app.route('/classify_file', methods=['POST'])
    @limiter.limit("10 per hour")
    def classify_file_route():
        if 'file' not in request.files:
            return jsonify({"error": "No file part in the request"}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400

        if not allowed_file(file.filename):
            return jsonify({"error": "File type not allowed"}), 400

        # Save the file temporarily
        file_extension = os.path.splitext(file.filename)[1].lower()
        temp_file_path = os.path.join('/data', 'temp_uploads', file.filename)
        os.makedirs(os.path.dirname(temp_file_path), exist_ok=True)
        file.save(temp_file_path)

        # Trim the file if necessary
        trimmed_file_path = temp_file_path
        if file_extension == '.pdf':
            reader = PdfReader(temp_file_path)
            if len(reader.pages) > LARGE_PAGE_COUNT:
                trimmed_file_path = temp_file_path.replace('.pdf', '_trimmed.pdf')
                trim_pdf_to_first_five_pages(temp_file_path, trimmed_file_path)
        elif file_extension == '.docx':
            document = Document(temp_file_path)
            if len(document.paragraphs) > LARGE_PAGE_COUNT:
                trimmed_file_path = temp_file_path.replace('.docx', '_trimmed.docx')
                trim_docx_to_first_five_paragraphs(temp_file_path, trimmed_file_path)

        # Proceed with normal processing
        # txt_filename = os.path.join('/data', OUTPUT_DIR, f"{os.path.splitext(file.filename)[0]}.txt")
        txt_filename = os.path.join('/tmp', OUTPUT_DIR, f"{os.path.splitext(file.filename)[0]}.txt") # Change to /tmp for Modal
        os.makedirs(os.path.dirname(txt_filename), exist_ok=True) 
        extracted_data = process_file(trimmed_file_path, txt_filename, file_extension)
        print("Finished OCR, extracted data: ", extracted_data)

        # Classify the document by file name
        document_type_by_name = classify_file_by_name(file)
        print("Finished classification by name, document type: ", document_type_by_name)

        # Classify the document using Llama 3
        if not extracted_data:
            document_type, reasoning = 'unknown file', 'Unable to extract data from the file'
        else:
            document_type, reasoning = classify_with_llama(extracted_data)
            print("Finished classification by content, document type: ", document_type)

        print("Finished processing file: ", file.filename)
        return jsonify({
            "document_type_from_llama": document_type,
            "llama_ai_classification_reasoning": reasoning,
            "classified_file_name": document_type_by_name,
            "filename": file.filename
        }), 200

    return app

# Local entry point for testing
@app.local_entrypoint()
def test_classify_with_llama():
    POSSIBLE_FILE_TYPES = [
    "Invoice",
    "Contract",
    "Resume",
    "Report",
    "Letter",
    "Email",
    "Research Paper",
    "Form",
    "Memo",
    "Receipt",
    "Bank Statement",
    "Medical Record",
    "Legal Document",
    "Other"
    ]
    model = Model()
    extracted_text = "Bank 1 Bank 1 of Testing Bank 1 - Confidential Statement | Page 2"
    prompt = (
        f"Classify the following document based on its content. Possible types are: {', '.join(POSSIBLE_FILE_TYPES)}.\n\n"
        f"Document Content:\n{extracted_text}\n\n"
        "Return the classification as follows in this exact format: 'Document Type: [Type]', Reasoning: [Reasoning]."
    )
    response = model.classify_with_llama.remote(prompt)
    print("Response from Llama model:")
    print(response)


# Entry point to run the app
if __name__ == '__main__':
    with app.run():
        flask_app()
